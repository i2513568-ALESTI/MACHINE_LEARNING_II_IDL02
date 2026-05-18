"""Apply visible grid borders to all tables in a Word document."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_table_borders(table, *, size: int = 4, color: str = "000000") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)

    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)
        borders.append(edge)
    tbl_pr.append(borders)


def fix_docx(path: Path) -> int:
    doc = Document(path)
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        set_table_borders(table)
    doc.save(path)
    return len(doc.tables)


def main() -> None:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("INFORME_ACADEMICO_CNN.docx")
    if not path.is_file():
        raise SystemExit(f"No existe: {path}")
    n = fix_docx(path)
    print(f"Bordes aplicados en {n} tablas: {path}")


if __name__ == "__main__":
    main()
