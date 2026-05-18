"""Regenera INFORME_ACADEMICO_CNN.docx desde el Markdown con tablas con bordes."""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "INFORME_ACADEMICO_CNN.md"
DOCX = ROOT / "INFORME_ACADEMICO_CNN.docx"
REF = ROOT / "pandoc-reference.docx"
FIX = Path(__file__).resolve().parent / "fix_docx_table_borders.py"


def ensure_reference_doc() -> None:
    if REF.exists():
        return
    from docx import Document

    doc = Document()
    doc.add_paragraph("Plantilla de referencia para Pandoc.")
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    for row in t.rows:
        for cell in row.cells:
            cell.text = "."
    doc.save(REF)


def main() -> None:
    ensure_reference_doc()
    cmd = [
        "pandoc",
        str(MD),
        "-o",
        str(DOCX),
        f"--resource-path={ROOT}",
        f"--reference-doc={REF}",
    ]
    subprocess.run(cmd, check=True, cwd=ROOT)
    subprocess.run([sys.executable, str(FIX), str(DOCX)], check=True, cwd=ROOT)


if __name__ == "__main__":
    main()
